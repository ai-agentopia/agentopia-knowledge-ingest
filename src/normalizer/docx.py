"""DOCX normalizer using python-docx."""

import logging

from normalizer.base import NormalizedDocument

logger = logging.getLogger(__name__)


def parse_docx(raw_bytes: bytes) -> NormalizedDocument:
    """Extract text and heading structure from DOCX bytes.

    Preserves heading hierarchy by inspecting paragraph styles.
    Tables are converted to tab-separated text.

    Raises:
        ImportError: If python-docx is not installed.
        Exception: If the file is not a valid DOCX.
    """
    import io
    from docx import Document

    doc = Document(io.BytesIO(raw_bytes))
    paragraphs: list[str] = []
    headings: list[str] = []
    has_tables = bool(doc.tables)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)

        # Word heading styles: Heading 1, Heading 2, ..., Title
        style_name = (para.style.name or "").lower()
        if "heading" in style_name or style_name == "title":
            headings.append(text)

    # Convert tables to tab-separated text and append
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = "\t".join(cells)
            if row_text.strip():
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)
    if not full_text.strip():
        raise ValueError("DOCX contains no extractable text")

    words = full_text.split()
    return NormalizedDocument(
        text=full_text,
        headings=headings,
        structure_hints={
            "has_tables": has_tables,
            "has_code_blocks": False,
            "heading_count": len(headings),
            "estimated_tokens": len(words),
        },
    )
